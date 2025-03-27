import os
import sys
import re
import argparse
import logging
from typing import Dict, List, Optional, Tuple, Union
from pathlib import Path

from rich.console import Console
from rich.panel import Panel
from rich.prompt import Prompt, Confirm
from rich.table import Table
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TimeElapsedColumn
from rich.syntax import Syntax
from rich.logging import RichHandler
from rich import print

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(message)s",
    datefmt="[%X]",
    handlers=[RichHandler(rich_tracebacks=True)]
)
logger = logging.getLogger("md_wizard")


class MarkdownParser:
    """Handles parsing of Markdown syntax into structured data."""
    
    def __init__(self):
        # Regex patterns for markdown elements
        self.patterns = {
            'header': re.compile(r'^(#{1,6})\s+(.+)$'),
            'bold': re.compile(r'\*\*(.*?)\*\*'),
            'italic': re.compile(r'\*(.*?)\*'),
            'code': re.compile(r'`(.*?)`'),
            'link': re.compile(r'\[(.*?)\]\((.*?)\)'),
            'list_item': re.compile(r'^\s*[-*+]\s+(.+)$'),
            'numbered_list': re.compile(r'^\s*(\d+)[\.\)]\s+(.+)$'),
            'blockquote': re.compile(r'^\s*>\s+(.+)$'),
            'horizontal_rule': re.compile(r'^(\s*[-*_]\s*){3,}$'),
            'image': re.compile(r'!\[(.*?)\]\((.*?)\)'),
        }
    
    def parse_line(self, line: str) -> dict:
        """
        Parse a single line of markdown text and return its structure.
        
        Args:
            line: A string containing markdown text
            
        Returns:
            A dictionary with the parsed structure
        """
        line = line.rstrip()
        
        # Check for headers
        header_match = self.patterns['header'].match(line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            return {'type': 'header', 'level': level, 'text': text}
        
        # Check for list items
        list_match = self.patterns['list_item'].match(line)
        if list_match:
            return {'type': 'list_item', 'text': list_match.group(1)}
        
        # Check for numbered list items
        numbered_match = self.patterns['numbered_list'].match(line)
        if numbered_match:
            return {
                'type': 'numbered_list', 
                'number': int(numbered_match.group(1)),
                'text': numbered_match.group(2)
            }
        
        # Check for blockquotes
        blockquote_match = self.patterns['blockquote'].match(line)
        if blockquote_match:
            return {'type': 'blockquote', 'text': blockquote_match.group(1)}
        
        # Check for horizontal rules
        if self.patterns['horizontal_rule'].match(line):
            return {'type': 'horizontal_rule'}
        
        # If no special format, it's a paragraph
        if line.strip():
            return {'type': 'paragraph', 'text': line}
        
        # Empty line
        return {'type': 'empty'}
    
    def find_inline_formats(self, text: str) -> List[dict]:
        """
        Find all inline formatting in a text string.
        
        Args:
            text: The text to parse for inline formatting
            
        Returns:
            A list of dictionaries with format information
        """
        formats = []
        
        # Find bold text
        for match in self.patterns['bold'].finditer(text):
            formats.append({
                'type': 'bold',
                'start': match.start(),
                'end': match.end(),
                'text': match.group(1),
                'original': match.group(0)
            })
        
        # Find italic text
        for match in self.patterns['italic'].finditer(text):
            # Skip if this is part of a bold format (e.g., **bold with *italic* inside**)
            if any(f['start'] <= match.start() < f['end'] for f in formats if f['type'] == 'bold'):
                continue
            formats.append({
                'type': 'italic',
                'start': match.start(),
                'end': match.end(),
                'text': match.group(1),
                'original': match.group(0)
            })
        
        # Find code snippets
        for match in self.patterns['code'].finditer(text):
            formats.append({
                'type': 'code',
                'start': match.start(),
                'end': match.end(),
                'text': match.group(1),
                'original': match.group(0)
            })
        
        # Find links
        for match in self.patterns['link'].finditer(text):
            formats.append({
                'type': 'link',
                'start': match.start(),
                'end': match.end(),
                'text': match.group(1),
                'url': match.group(2),
                'original': match.group(0)
            })
        
        # Find images
        for match in self.patterns['image'].finditer(text):
            formats.append({
                'type': 'image',
                'start': match.start(),
                'end': match.end(),
                'alt_text': match.group(1),
                'url': match.group(2),
                'original': match.group(0)
            })
        
        # Sort formats by start position
        return sorted(formats, key=lambda x: x['start'])


class DocxStyler:
    """Handles styling and formatting of the DOCX document."""
    
    def __init__(self, document: Document, theme: str = "default"):
        self.document = document
        self.theme = theme
        self.setup_styles()
    
    def setup_styles(self):
        """Set up custom styles for the document based on the selected theme."""
        styles = self.document.styles
        
        # Title style (Heading 1)
        title_style = styles.add_style('MDTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.base_style = styles['Heading 1']
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        if self.theme == "professional":
            title_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Subtitle style (Heading 2)
        subtitle_style = styles.add_style('MDSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.base_style = styles['Heading 2']
        subtitle_style.font.size = Pt(18)
        subtitle_style.font.bold = True
        if self.theme == "professional":
            subtitle_style.font.color.rgb = RGBColor(0, 102, 153)  # Medium blue
        
        # Heading 3 style
        h3_style = styles.add_style('MDHeading3', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.base_style = styles['Heading 3']
        h3_style.font.size = Pt(14)
        h3_style.font.bold = True
        if self.theme == "professional":
            h3_style.font.color.rgb = RGBColor(0, 128, 192)  # Light blue
        
        # Blockquote style
        blockquote_style = styles.add_style('MDBlockquote', WD_STYLE_TYPE.PARAGRAPH)
        blockquote_style.font.italic = True
        blockquote_style.paragraph_format.left_indent = Inches(0.5)
        if self.theme == "professional":
            blockquote_style.font.color.rgb = RGBColor(102, 102, 102)  # Gray
        
        # Code style
        code_style = styles.add_style('MDCode', WD_STYLE_TYPE.CHARACTER)
        code_style.font.name = 'Courier New'
        if self.theme == "professional":
            code_style.font.color.rgb = RGBColor(153, 51, 0)  # Rust color
        
        # List item style
        list_style = styles.add_style('MDListItem', WD_STYLE_TYPE.PARAGRAPH)
        list_style.paragraph_format.left_indent = Inches(0.25)
        list_style.paragraph_format.first_line_indent = Inches(-0.25)
    
    def add_hyperlink(self, paragraph, text: str, url: str):
        """Add a hyperlink to a paragraph."""
        # This gets access to the document's XML
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        # Create the hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # Create a new run element
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Style the hyperlink
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        # Add underline
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        
        # Add text to the run
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        
        return hyperlink


class MarkdownConverter:
    def __init__(self, theme: str = "default"):
        self.console = Console()
        self.converter_name = "MD Wizard 🧙‍♂️"
        self.parser = MarkdownParser()
        self.theme = theme
    
    def show_banner(self):
        """Display an attractive banner for the application."""
        self.console.print(Panel(
            f"[bold blue]{self.converter_name}[/bold blue]\n[dim]Advanced Markdown to DOCX Converter[/dim]",
            border_style="blue",
            expand=False,
            padding=(1, 10)
        ), justify="center")
    
    def fix_encoding(self, text: str) -> str:
        """Fix common encoding issues in text."""
        encoding_map = {
            'Ã³': 'ó', 'Ã¡': 'á', 'Ã©': 'é', 'Ã­': 'í', 'Ãº': 'ú',
            'Ã±': 'ñ', 'Ãƒ': 'í', 'Ã‰': 'É', 'Ã�': 'Á', 'Ã©': 'é',
            'Ã³': 'ó', 'Ãº': 'ú'
        }
        for key, value in encoding_map.items():
            text = text.replace(key, value)
        return text
    
    def select_markdown_file(self) -> str:
        """Interactive file selection with rich validation."""
        while True:
            file_path = Prompt.ask("[green]Enter the path to your Markdown file[/green]")
            if not file_path:
                self.console.print("[red]❌ File path cannot be empty![/red]")
                continue
            
            path = Path(file_path)
            if not path.exists():
                self.console.print(f"[red]❌ File {file_path} does not exist![/red]")
                continue
                
            if not path.is_file():
                self.console.print(f"[red]❌ {file_path} is not a file![/red]")
                continue
                
            if path.suffix.lower() != '.md':
                if not Confirm.ask(f"[yellow]⚠️ {file_path} doesn't have a .md extension. Are you sure it's a Markdown file?[/yellow]"):
                    continue
            
            return str(path)
    
    def select_output_file(self, input_file: str) -> str:
        """Interactive output file selection with smart defaults."""
        default_output = os.path.splitext(input_file)[0] + '.docx'
        output_file = Prompt.ask(
            "[green]Enter output DOCX file path[/green]", 
            default=default_output
        )
        
        # Check if file exists and confirm overwrite
        if os.path.exists(output_file):
            if not Confirm.ask(f"[yellow]⚠️ File {output_file} already exists. Overwrite?[/yellow]"):
                return self.select_output_file(input_file)  # Ask again
        
        return output_file
    
    def select_theme(self) -> str:
        """Allow user to select a document theme."""
        table = Table(title="Available Themes")
        table.add_column("Theme", style="cyan")
        table.add_column("Description")
        
        table.add_row("default", "Standard formatting with minimal styling")
        table.add_row("professional", "Business-oriented with blue headings")
        
        self.console.print(table)
        
        theme = Prompt.ask(
            "[green]Select a theme[/green]",
            choices=["default", "professional"],
            default="default"
        )
        
        return theme
    
    def apply_inline_formatting(self, paragraph, text: str, formats: List[dict]):
        """
        Apply inline formatting to a paragraph based on parsed formats.
        
        Args:
            paragraph: The docx paragraph to add text to
            text: The original text
            formats: List of format dictionaries from the parser
        """
        if not formats:
            # No special formatting, add the whole text
            paragraph.add_run(text)
            return
        
        last_end = 0
        for fmt in formats:
            # Add text before this format
            if fmt['start'] > last_end:
                paragraph.add_run(text[last_end:fmt['start']])
            
            # Handle different format types
            if fmt['type'] == 'bold':
                run = paragraph.add_run(fmt['text'])
                run.bold = True
            elif fmt['type'] == 'italic':
                run = paragraph.add_run(fmt['text'])
                run.italic = True
            elif fmt['type'] == 'code':
                run = paragraph.add_run(fmt['text'])
                run.style = 'MDCode'
            elif fmt['type'] == 'link':
                # Add hyperlink
                styler = DocxStyler(paragraph.part.document)
                styler.add_hyperlink(paragraph, fmt['text'], fmt['url'])
            elif fmt['type'] == 'image':
                # Images are handled separately
                pass
            
            last_end = fmt['end']
        
        # Add any remaining text after the last format
        if last_end < len(text):
            paragraph.add_run(text[last_end:])
    
    def markdown_to_docx(self, input_file: str, output_file: str):
        """
        Convert markdown file to DOCX with rich progress tracking.
        
        Args:
            input_file: Path to the markdown file
            output_file: Path where the DOCX file will be saved
        """
        try:
            progress = Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                TimeElapsedColumn()
            )
            
            with progress:
                # Read markdown file
                read_task = progress.add_task("[cyan]Reading Markdown...", total=100)
                try:
                    with open(input_file, 'r', encoding='utf-8') as f:
                        markdown_text = f.read()
                except UnicodeDecodeError:
                    # Try with different encodings if utf-8 fails
                    encodings = ['latin-1', 'iso-8859-1', 'cp1252']
                    for encoding in encodings:
                        try:
                            with open(input_file, 'r', encoding=encoding) as f:
                                markdown_text = f.read()
                            logger.info(f"Successfully read file with {encoding} encoding")
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        raise ValueError("Could not decode the file with any supported encoding")
                
                # Fix encoding issues
                markdown_text = self.fix_encoding(markdown_text)
                progress.update(read_task, completed=100)
                
                # Parse markdown
                parse_task = progress.add_task("[yellow]Parsing Markdown...", total=100)
                lines = markdown_text.split('\n')
                parsed_lines = []
                
                for i, line in enumerate(lines):
                    parsed = self.parser.parse_line(line)
                    if parsed['type'] in ['paragraph', 'header', 'list_item', 'numbered_list', 'blockquote']:
                        # Find inline formatting
                        if 'text' in parsed:
                            formats = self.parser.find_inline_formats(parsed['text'])
                            parsed['formats'] = formats
                    
                    parsed_lines.append(parsed)
                    progress.update(parse_task, completed=int((i+1)/len(lines)*100))
                
                # Create document
                convert_task = progress.add_task("[green]Converting to DOCX...", total=100)
                doc = Document()
                styler = DocxStyler(doc, self.theme)
                
                # Process parsed lines
                list_level = 0
                in_list = False
                
                for i, parsed in enumerate(parsed_lines):
                    progress.update(convert_task, completed=int((i+1)/len(parsed_lines)*100))
                    
                    if parsed['type'] == 'empty':
                        # Add empty paragraph for spacing
                        if i > 0 and parsed_lines[i-1]['type'] != 'empty':
                            doc.add_paragraph()
                        continue
                    
                    elif parsed['type'] == 'header':
                        # Add header with appropriate style
                        level = parsed['level']
                        if level == 1:
                            p = doc.add_paragraph(parsed['text'], style='MDTitle')
                        elif level == 2:
                            p = doc.add_paragraph(parsed['text'], style='MDSubtitle')
                        elif level == 3:
                            p = doc.add_paragraph(parsed['text'], style='MDHeading3')
                        else:
                            p = doc.add_paragraph(parsed['text'])
                            p.style = f'Heading {min(level, 9)}'
                    
                    elif parsed['type'] == 'paragraph':
                        # Add paragraph with inline formatting
                        p = doc.add_paragraph()
                        self.apply_inline_formatting(p, parsed['text'], parsed.get('formats', []))
                    
                    elif parsed['type'] == 'list_item':
                        # Handle unordered list items
                        if not in_list:
                            in_list = True
                        
                        p = doc.add_paragraph(style='MDListItem')
                        p.paragraph_format.left_indent = Inches(0.25 * list_level)
                        
                        # Add bullet
                        run = p.add_run("• ")
                        run.bold = True
                        
                        # Add content with formatting
                        self.apply_inline_formatting(p, parsed['text'], parsed.get('formats', []))
                    
                    elif parsed['type'] == 'numbered_list':
                        # Handle ordered list items
                        if not in_list:
                            in_list = True
                        
                        p = doc.add_paragraph(style='MDListItem')
                        p.paragraph_format.left_indent = Inches(0.25 * list_level)
                        
                        # Add number
                        run = p.add_run(f"{parsed['number']}. ")
                        run.bold = True
                        
                        # Add content with formatting
                        self.apply_inline_formatting(p, parsed['text'], parsed.get('formats', []))
                    
                    elif parsed['type'] == 'blockquote':
                        # Add blockquote
                        p = doc.add_paragraph(style='MDBlockquote')
                        self.apply_inline_formatting(p, parsed['text'], parsed.get('formats', []))
                    
                    elif parsed['type'] == 'horizontal_rule':
                        # Add horizontal line
                        p = doc.add_paragraph()
                        p.paragraph_format.border_bottom = True
                        p.paragraph_format.space_after = Pt(12)
                    
                    # Check if we're exiting a list
                    if in_list and i < len(parsed_lines) - 1:
                        next_parsed = parsed_lines[i+1]
                        if next_parsed['type'] not in ['list_item', 'numbered_list']:
                            in_list = False
                            list_level = 0
                
                # Save document
                save_task = progress.add_task("[magenta]Saving document...", total=100)
                doc.save(output_file)
                progress.update(save_task, completed=100)
            
            self.console.print(Panel(
                f"[bold green]✅ Document successfully saved to {output_file}[/bold green]",
                border_style="green"
            ))
            
        except Exception as e:
            logger.error(f"Conversion failed: {str(e)}", exc_info=True)
            self.console.print(f"[bold red]❌ Conversion failed: {str(e)}[/bold red]")
    
    def run(self):
        """Main application flow."""
        self.show_banner()
        
        try:
            # Parse command line arguments if provided
            parser = argparse.ArgumentParser(description='Convert Markdown to DOCX')
            parser.add_argument('-i', '--input', help='Input markdown file')
            parser.add_argument('-o', '--output', help='Output DOCX file')
            parser.add_argument('-t', '--theme', choices=['default', 'professional'], 
                                default='default', help='Document theme')
            
            args, unknown = parser.parse_known_args()
            
            # Set theme if provided
            if args.theme:
                self.theme = args.theme
            else:
                self.theme = self.select_theme()
            
            # Get input file
            if args.input and os.path.exists(args.input):
                input_file = args.input
            else:
                input_file = self.select_markdown_file()
            
            # Get output file
            if args.output:
                output_file = args.output
            else:
                output_file = self.select_output_file(input_file)
            
            # Show preview of markdown content
            try:
                with open(input_file, 'r', encoding='utf-8') as f:
                    markdown_preview = f.read(500)  # First 500 chars
            except UnicodeDecodeError:
                # Try with different encoding
                with open(input_file, 'r', encoding='latin-1') as f:
                    markdown_preview = f.read(500)
            
            markdown_preview = self.fix_encoding(markdown_preview)
            
            self.console.print("\n[bold]Markdown Preview:[/bold]")
            syntax = Syntax(markdown_preview + "..." if len(markdown_preview) >= 500 else markdown_preview, 
                           "markdown", theme="monokai")
            self.console.print(Panel(syntax, border_style="dim"))
            
            if args.input and args.output:
                # If both input and output are provided via command line, proceed without confirmation
                self.markdown_to_docx(input_file, output_file)
            else:
                if Confirm.ask("[yellow]Proceed with conversion?[/yellow]"):
                    self.markdown_to_docx(input_file, output_file)
                else:
                    self.console.print("[red]Conversion cancelled.[/red]")
                    
        except KeyboardInterrupt:
            self.console.print("\n[red]Operation cancelled by user.[/red]")
            return 1
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}", exc_info=True)
            self.console.print(f"[bold red]❌ An unexpected error occurred: {str(e)}[/bold red]")
            return 1
            
        return 0


def main():
    """Entry point for the application."""
    try:
        converter = MarkdownConverter()
        return converter.run()
    except Exception as e:
        logger.critical(f"Fatal error: {str(e)}", exc_info=True)
        Console().print(f"[bold red]FATAL ERROR: {str(e)}[/bold red]")
        return 1


if __name__ == "__main__":
    sys.exit(main())